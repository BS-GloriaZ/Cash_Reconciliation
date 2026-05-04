"""Generate RecX Technical Specification as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text=""):
    return doc.add_paragraph(text)

def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "AAAAAA")
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def note_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x7B, 0x36, 0x00)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF3CD")
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF3FB")
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("RecX — Cash Reconciliation Pipeline", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Technical Specification")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].bold = True
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. DATA INPUTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Data Inputs")

# 1.1 Tradar
heading("1.1  Tradar (Internal Portfolio System)", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*Cash Flow since One Month ago - all funds.csv"],
        ["Format", "CSV, UTF-8-BOM"],
        ["Structure", "Multi-section: Row 1 = title with date range (DD Mon YY – DD Mon YY); Row 2 = header; subsequent rows interspersed with Fund:, Account:, Ccy: label rows"],
        ["Key columns", "Trade, Type, Amount, Sedol, Isin, Security Type, Description, Price, Date, Settles, Cashflow, Cumulative Cashflow, Balance, Notes"],
    ],
    col_widths=[1.5, 5.0],
)
para("The file covers all funds in one export. Fund/Account/CCY context is inherited from preceding label rows.")

# 1.2 Citi Balances
heading("1.2  Citi — Cash Balances", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "DOD_CASH_BALANCES_V1_C30_*.csv"],
        ["Source folder", r"W:\Ops\Citi\Rec Files\Archive"],
        ["Key columns", "Account ID, Account Name, Currency Code, Close of Business Date, Ledger Balance, Available Balance, Opening Balance, As of Date, Balance Update Timestamp"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.3 Citi HI
heading("1.3  Citi — High Interest (HI) Positions", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["Files", "BSFFTMOUT_Positions_All_*.CSV (one file per settlement date)"],
        ["Source folder", r"W:\Ops\Citi\Rec Files"],
        ["Filename date", "YYMMDD embedded in stem"],
        ["Key columns", "Account ID (= SKAC), Sec ID (= security/fund account), Available Position"],
        ["Note", "If two files exist for the same date (AM + same-day), only the one with the latest modification time is loaded"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.4 Citi Txns
heading("1.4  Citi — Cash Transactions", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "DOD_CASH_TRANSACTIONS_V1_C30_*.csv (latest file only — contains full history)"],
        ["Source folder", r"W:\Ops\Citi\Rec Files\Archive"],
        ["Key columns", "Account ID, Currency Code, Amount, Close of Business Date, Contractual Settlement Date, Transaction Description, SEDOL, ISIN, Issue Description"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.5 BNP Balances
heading("1.5  BNP — Cash Balances", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*GPBCash*.csv"],
        ["Source folder", r"W:\Ops\BNP\Rec Files"],
        ["Key columns", "ProcessDate, AsOfDate (YYYYMMDD), AccountCode, AccountName, CurrencyCode, SettleDateBalanceLocal"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.6 BNP Txns
heading("1.6  BNP — Cash Transactions", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*99X.CashLedgerSD.csv"],
        ["Source folder", r"W:\Ops\BNP\Rec Files"],
        ["Key columns", "AsOfDate, Settle Date (YYYYMMDD integer), ID (account), Code (ISO currency), Net Amount, Activity (description), SEDOL, ISIN, Security Description, Previous Local Opening Balance, Local Closing Balance"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.7 BNP NZ Balances
heading("1.7  BNP NZ — Cash Balances", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*Bal_cash*.csv"],
        ["Source folder", r"W:\Ops\BNP NZ\Rec Files\Recon cleaned"],
        ["Key columns", "Account ID, Account Name, Account Base Currency Code, Contractual Settlement Date, Opening Balance, Closing Balance, Created Timestamp"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.8 BNP NZ Txns
heading("1.8  BNP NZ — Cash Transactions", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*_BNPNZ_Custody_Cash.csv"],
        ["Source folder", r"W:\Ops\BNP NZ\Rec Files\Recon cleaned"],
        ["Key columns", "Account ID, Contractual Settlement Date (or Entry Date), Value Date, Account Base Currency Code, Amount, Cash Transaction Type, SEDOL, ISIN"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.9 BNP Margin PDF
heading("1.9  BNP Margin PDF (BBUS Special)", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "*DAILY_STAT.pdf"],
        ["Source folder", r"W:\Ops\BNP\Rec Files"],
        ["Extracted values", "Initial Margin Requirement USD, Variation Margin row (Previous Unrealized Trading/Hedging, Unrealized Trading/Hedging, Mark to Market)"],
        ["Derived adjustment", "Total PDF Adjustment = Initial Margin + (−Unrealized Trading)"],
    ],
    col_widths=[1.5, 5.0],
)

# 1.10 Mapping
heading("1.10  Account Mapping File", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "mapping.xlsx (all sheets concatenated)"],
        ["Key columns", "Fund, Portfolio, Account ID, Currency Code, Custody, Tradar_Account, Fund Type, SKAC"],
        ["Column aliases", 'Flexible — resolved via config (e.g. "Fund Code", "Mapping Fund", "FUND NAME (F59/72)" all resolve to Fund)'],
    ],
    col_widths=[1.5, 5.0],
)

# 1.11 T-Log
heading("1.11  T-Log Flow Reference File", 2)
add_table(
    ["Attribute", "Detail"],
    [
        ["File", "uploaded_tlogdata_YYYYMMDD HHMMSS.csv"],
        ["Source folder", r"W:\Ops\Templates and Trades for RBC\_tlog_data\Processed"],
        ["Local subdir", "data/input/tlog"],
        ["Retention", "All files within lookback window (multi-file)"],
    ],
    col_widths=[1.5, 5.0],
)
para("Key columns:")
add_table(
    ["Column", "Description"],
    [
        ["fund", "Fund code — matches Mapping Fund"],
        ["units_order", "Units ordered. Positive = Creation, Negative = Redemption"],
        ["ccy", "Settlement currency"],
        ["trade_date", "Trade date"],
        ["settle_date", "Expected cash settlement date"],
        ["gross_aud", "Gross consideration in AUD"],
        ["gross_local", "Gross consideration in local currency (used when ccy ≠ AUD)"],
        ["fixed_fee", "Fixed fee amount"],
        ["var_fee", "Variable fee amount (may be zero/blank)"],
        ["net", "Net transfer amount (final, when NAV is known)"],
        ["nav_est", "Estimated NAV — populated when final NAV not yet available at settlement"],
        ["est_net", "Estimated net transfer amount based on nav_est"],
        ["trueup", "True-up transfer amount — corrective transfer arriving once actual NAV is finalised"],
    ],
    col_widths=[1.5, 5.0],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. DATA CLEANING
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Data Cleaning Rules")

heading("2.1  Tradar", 2)
for item in [
    "File parsed raw (no pandas header inference); Row 1 is the title (used only for date range extraction); Row 2 is the column header.",
    "Blank column names are back-filled from a fixed positional template.",
    "Duplicate column names are de-duplicated with a numeric suffix.",
    "Fund:, Account:, Ccy: label rows are stripped out; their values are propagated forward to all subsequent data rows until the next label.",
    "Dates (Date, Settles) parsed as %d %b %Y.",
    "Numeric fields (Amount, Cashflow, Cumulative Cashflow, Balance, Price) stripped of commas/whitespace and converted to float.",
    "A boolean flag Is Opening Balance is set where Type == 'opening balance' (case-insensitive).",
    "Empty rows (all non-context fields blank) are dropped.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.2  Citi Balances", 2)
for item in [
    "Sorted by (Date, Account ID, Currency Code, Balance Update Timestamp, As of Date).",
    "Deduplicated: for each (Date, Account ID, Currency Code), only the last row is kept (most recent timestamp wins).",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.3  Citi HI Positions", 2)
for item in [
    "One file per settlement date; AM vs same-day duplicates resolved by keeping the file with the latest modification time.",
    "Files outside the lookback window are excluded.",
    "Source Type = 'CITI_HI' flag added to distinguish from normal Citi balances.",
    "No currency code in source — currency is filled from the mapping file at reconciliation time.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.4  Citi Transactions", 2)
for item in [
    "Only the latest file is loaded (it contains full transaction history; loading multiple files would double-count rows).",
    "Security ID = SEDOL if non-blank, otherwise ISIN.",
    "Rows with null COB Date or null Amount are dropped.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.5  BNP Balances", 2)
for item in [
    "AsOfDate parsed as YYYYMMDD; ProcessDate as YYYYMMDD HH:MM:SS.",
    "Deduplicated: keep last row per (AsOfDate, AccountCode, CurrencyCode).",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.6  BNP Transactions", 2)
for item in [
    "Balance summary rows (Amount = 0 or null) are separated and used to derive Opening Balance / Closing Balance, then joined back onto transaction rows by (COB Date, Account ID, Currency Code).",
    "Settle Date parsed from YYYYMMDD integer string.",
    "Security ID = SEDOL if non-blank, otherwise ISIN.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.7  BNP NZ Balances", 2)
for item in [
    "Contractual Settlement Date parsed as date; Created Timestamp as datetime.",
    "Deduplicated: keep last row per (Settlement Date, Account ID, Currency Code).",
    "Closing Balance used as the ledger balance.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.8  BNP NZ Transactions", 2)
for item in [
    "COB Date sourced from Contractual Settlement Date if present, otherwise Entry Date.",
    "Settle Date from Value Date.",
    "Description from Cash Transaction Type.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("2.9  Mapping", 2)
for item in [
    "All sheets concatenated.",
    "Column names resolved via aliases (flexible headers).",
    "Fund uppercased; Account ID and SKAC normalised (strip, uppercase).",
    "Rows where Fund is blank are marked In Scope = False.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MATCHING / RECONCILIATION LOGIC
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Matching / Reconciliation Logic")

heading("3.1  Account Matching (Source → Mapping)", 2)
add_table(
    ["Custodian", "Match Keys"],
    [
        ["Citi", "(Source Account ID, Currency Code) → (Account ID, Currency Code)"],
        ["Citi HI", "(Source SKAC, Source Account ID) → (SKAC, Account ID)"],
        ["BNP", "Source Account ID → Account ID; currency validated if mapping has currency"],
        ["BNP NZ", "Derives fund code by stripping currency suffix from account ID (e.g. BBUSNZD → BBUS); matches on derived fund + currency"],
    ],
    col_widths=[1.5, 5.0],
)
para("Tradar_Account in the mapping supports wildcard suffixes: RBC* matches any Tradar account starting with RBC.")

heading("3.2  Tradar Daily Balance Computation", 2)
for item in [
    "Filter Tradar rows to mapped accounts only.",
    "Extract opening balance per (Fund, CCY, Account) — the row where Type = 'Opening Balance'.",
    "For settled transactions: take the last transaction per (Fund, CCY, Account, Settles) — its Balance field is the end-of-day balance for that settle date.",
    "Build a date grid: every (Fund, CCY, Account) × every business date in the lookback window.",
    "Use merge_asof (backward) to find, for each (group, COB date), the balance as of the last settle date ≤ COB date.",
    "Where no settled transaction exists, fall back to the opening balance.",
]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(item)
para()

heading("3.3  Balance Reconciliation (Normal Rec)", 2)
for item in [
    "For each (Date, Fund, Currency Code, Tradar_Account):",
    "Variance = Source Ledger Balance − Tradar Balance",
    "Status = Matched if round(|Variance|, 2) == 0, else Break",
    "Only in-scope, mapped rows are included",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

heading("3.4  High Interest Reconciliation", 2)
para("Same logic as 3.3 but using CITI_HI source rows. Match on (SKAC, Sec ID) → (SKAC, Account ID) in mapping. Currency filled from mapping.")

heading("3.5  BBUS BNP Special Reconciliation", 2)
for item in [
    "Source side: Sum all BNP USD account balances mapping to BBUS as of the latest BNP date ≤ run date.",
    "PDF adjustment: Add Total PDF Adjustment USD (= Initial Margin − Unrealized Trading P&L).",
    "Adjusted Source = BNP Cash + PDF Adjustment.",
    "Tradar side: Sum balances across three specific accounts (RBC-CS, FUT USD, BTMU 11am) in USD.",
    "Variance = Adjusted Source − Tradar Total",
]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(item)
para()

heading("3.6  Transaction Reconciliation (Waterfall / Break Analysis)", 2)
para("For a given (Fund, Custody, Currency) break, Tradar settled transactions and custodian transaction records are matched using a multi-pass algorithm.")
para()
para("Security ID selection (Tradar side):")
add_table(
    ["Security Type", "Identifier Used"],
    [
        ["Government Bond, MTN Bond, Corporate Bond, Fixed Income, FRN, Floating Rate Note, Term Deposit", "ISIN"],
        ["Equity, Common Stock, Ordinary Share, Preference Share, REIT", "SEDOL"],
        ["Exchrate, FX, Exchange Rate", "None (routed to FX pass)"],
        ["Unknown", "SEDOL first, ISIN as fallback"],
    ],
    col_widths=[3.5, 3.0],
)

p = para()
p.add_run("Pass 1 — Dividend / Coupon grouping").bold = True
for item in [
    "Tradar rows where Type starts with div, cpn, or contains coupon.",
    "Grouped by (Security ID, Settle Date) and matched against custodian rows with dividend/interest/income descriptions at the same key. Comparison on group sums.",
    "Sub-pass B: single custodian row with same security ID and exact settle date (trade-type rows excluded).",
    "Sub-pass C: fallback — same security ID, same sign, settle within 20 calendar days.",
    "Trade-type custodian exclusion keywords: receipt vs payt, rvp, security purchases, purchase of securities, o/s purchase, sec-del vs payt, dvp, delivery vs payt, sale of securities, securities sale, o/s sale, contractual settlement.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

p = para()
p.add_run("Pass 2a — Security trades").bold = True
for item in [
    "Tradar Type = Purchase / Sell / Sale with a security ID.",
    "Matched by (Security ID, Settle Date) with matching trade description type.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

p = para()
p.add_run("Pass 2b — FX / cash trades").bold = True
for item in [
    "Tradar Type = Purchase / Sell / Sale with no security ID.",
    "Matched by settle date + amount within ±0.01.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

p = para()
p.add_run("Pass 3 — Generic 1-to-1 by exact amount").bold = True
for item in [
    "Hash-based O(n+m) lookup by integer cents. Greedy first-match.",
    "FT INTERNAL TRANSFER rows are excluded from this pass (reserved for the tlog-based flow cash pass).",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

para("Type Category classification:")
add_table(
    ["Category", "Condition"],
    [
        ["FX", "Security Type is Exchrate/FX, or purchase/sell with no security ID"],
        ["Dividend/Coupon", "Type starts with div or cpn, or descriptions contain dividend/coupon"],
        ["Interest", "Description contains 'interest'"],
        ["Tax", "Description contains tax/withholding/wht"],
        ["Trade", "Purchase/Sell with security ID, or custodian trade description"],
        ["Flow Cash", "Type is MiscIncome/MiscExpense, or description contains subscription/redemption/transfer"],
        ["Other", "Everything else"],
    ],
    col_widths=[1.8, 4.7],
)

heading("3.7  Flow Cash Matching — tlog-Based", 2)
para("This pass uses the T-Log reference file (Section 1.11) as the source of truth for creation and redemption cash flows. Each tlog row defines exactly which Tradar entries and custody transfers are expected.")
para()
note_block(
    "Important — Citi FT INTERNAL TRANSFER scope\n\n"
    "FT INTERNAL TRANSFER entries on the Citi custodian side are not exclusively flow-related. They also appear for "
    "ad-hoc inter-account transfers, cash movements, and other operational activity unrelated to creations or redemptions. "
    "Matching must therefore be driven strictly by the tlog reference file — a custodian FT INTERNAL TRANSFER entry should "
    "only be matched to a flow if a corresponding tlog row exists with a matching fund, currency, and expected transfer amount. "
    "Any FT INTERNAL TRANSFER entry that cannot be traced to a tlog row must not be automatically classified as a flow and "
    "should remain unmatched (Custodian Only) for manual review. "
    "This rule applies to Citi-custodied funds only; BNP and BNP NZ do not use the FT INTERNAL TRANSFER description for this purpose."
)
para()

para("Flow direction:")
add_table(
    ["units_order", "Flow type", "Expected Tradar type"],
    [
        ["Positive", "Creation", "MiscIncome"],
        ["Negative", "Redemption", "MiscExpense"],
    ],
    col_widths=[1.5, 2.0, 3.0],
)

p = para()
p.add_run("Step 1 — Tradar entry matching").bold = True
para("For each tlog row, locate the corresponding Tradar entries by (fund, settle_date):")
add_table(
    ["tlog column", "Expected Tradar entry"],
    [
        ["gross_aud (if ccy = AUD) or gross_local (if ccy ≠ AUD)", "1st MiscIncome/MiscExpense — amount must match exactly"],
        ["var_fee (if non-zero/non-blank)", "2nd MiscIncome/MiscExpense — amount must match exactly"],
    ],
    col_widths=[3.0, 3.5],
)
for item in [
    "If var_fee is zero or blank → one Tradar entry expected.",
    "If var_fee is populated → two Tradar entries expected.",
    "Status: Matched if all expected entries found with correct amounts; Break if amounts differ; Tradar Only if entries absent.",
]:
    doc.add_paragraph(item, style="List Bullet")
para()

p = para()
p.add_run("Step 2 — Custody transfer matching (Citi only)").bold = True
para("Only FT INTERNAL TRANSFER entries that can be linked back to a tlog row are eligible for matching. Entries with no corresponding tlog row are left as Custodian Only regardless of amount.")
para()

p = para()
p.add_run("Sub-case A — Final NAV (nav_est is blank)").italic = True
para("Expected custody transfer = net − fixed_fee")
para("Search Citi custodian FT INTERNAL TRANSFER entries for the matching fund/currency within ±5 calendar days of settle_date.")
add_table(
    ["Result", "Status"],
    [
        ["Custody amount found and matches expected", "Matched"],
        ["Custody amount found but differs", "Break — amount diff"],
        ["No custody entry found", "Break — awaiting transfer"],
    ],
    col_widths=[3.5, 3.0],
)

p = para()
p.add_run("Sub-case B — Estimated NAV (nav_est is populated)").italic = True
para("Settlement occurs before final NAV is known. Two custody transfers are expected:")
for item in [
    "Initial estimated transfer = est_net − fixed_fee  (expected around settle_date, ±5 calendar days)",
    "True-up transfer = trueup column value  (expected within 20 business days of settle_date once actual NAV is finalised)",
]:
    doc.add_paragraph(item, style="List Number")
para()
add_table(
    ["Condition", "Status"],
    [
        ["Both initial and true-up found with correct amounts", "Matched"],
        ["Initial found; true-up not yet arrived", "Break — awaiting true-up"],
        ["Initial found; true-up arrived but wrong amount", "Break — true-up amount diff"],
        ["Initial transfer not found", "Break — awaiting initial transfer"],
    ],
    col_widths=[3.5, 3.0],
)

para("Summary of expected entries per tlog row:")
code_block(
    "tlog row  (creation, AUD, var_fee present, nav_est present)\n"
    "|\n"
    "├─ Tradar side\n"
    "|   ├─ MiscIncome   amount = gross_aud      <- gross consideration\n"
    "|   └─ MiscIncome   amount = var_fee        <- variable fee\n"
    "|\n"
    "└─ Custody side  (Citi FT INTERNAL TRANSFER only)\n"
    "    |   (must be traceable to this tlog row; unlinked transfers excluded)\n"
    "    ├─ Initial transfer   amount = est_net - fixed_fee   (around settle_date)\n"
    "    └─ True-up transfer   amount = trueup                (within 20 business days)"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. DATA OUTPUTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Data Outputs")

heading("4.1  Normal Rec Detail", 2)
para("One row per (Date, Fund, Custody, Currency, Tradar Account).")
code_block(
    "Date | Custody | Fund Type | Source Account ID | Currency Code | Portfolio | Mapping Fund\n"
    "| Tradar_Account | Source Account Name | Source Ledger Balance | Source Available Balance\n"
    "| Source Opening Balance | Tradar Opening Balance | Tradar Settles Used | Tradar Balance\n"
    "| Variance | Abs Variance | Local Variance | BPS Impact | Materiality | Status"
)

heading("4.2  HI Rec Detail", 2)
para("One row per (Settlement Date, Fund, Currency, Tradar Account).")
code_block(
    "Target Date | Custody | Fund Type | Currency Code | Source SKAC | Source Account ID\n"
    "| Portfolio | Mapping Fund | Tradar_Account | Source Ledger Balance | Tradar Balance\n"
    "| Variance | Abs Variance | Local Variance | BPS Impact | Materiality | Status"
)

heading("4.3  BBUS BNP Special Detail", 2)
para("One row per run date.")
code_block(
    "Date | Fund | Custody | Currency Code | Portfolio | Mapping Fund | Fund Type\n"
    "| Source Account ID | BNP Cash Balance USD | Initial Margin Requirement USD\n"
    "| Unrealized Trading USD | Opposite Unrealized Trading USD | Total PDF Adjustment USD\n"
    "| Source Ledger Balance | Tradar Accounts Used | Tradar Balance | Variance | Abs Variance | Status"
)

heading("4.4  Transaction Rec Output", 2)
para("One row per matched group or unmatched transaction.")
code_block(
    "Counterparty (Both/Tradar/Custody) | Tradar Account | Type Category | Type\n"
    "| Security ID | Security Name | Tradar Settle | Custodian Settle\n"
    "| Tradar Amount | Custodian Amount | Direction | Description | Cust Description\n"
    "| Cust Type | Amount Diff | Settle Date Diff | Status"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Key Configuration Parameters")
add_table(
    ["Parameter", "Default", "Description"],
    [
        ["lookback_days", "60", "Business days of history to reconcile"],
        ["mapped_funds_only", "true", "Exclude accounts not in the mapping file"],
        ["source_balance_column", "Source Ledger Balance", "Custodian balance field used for comparison"],
        ["status_round_dp", "2", "Decimal places for variance rounding"],
        ["status_match_value", "0", "Variance threshold for Matched status"],
    ],
    col_widths=[2.2, 1.2, 3.1],
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/gloriazuo/Cash-Reconciliation_V2/RecX_Technical_Specification.docx"
doc.save(out)
print(f"Saved: {out}")
